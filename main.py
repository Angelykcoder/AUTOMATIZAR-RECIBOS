import sys
import calendar
import shutil
import tempfile
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer

# =====================================================================
# MODULE: UTILS & LOGIC
# =====================================================================

def obtener_ultimo_dia(anio: int, mes: int) -> int:
    """
    Calcula el último día de un mes y año específicos utilizando calendar.monthrange.
    """
    _, ultimo_dia = calendar.monthrange(anio, mes)
    return ultimo_dia

def formatear_moneda(monto: float) -> str:
    """
    Formatea un número flotante al formato de moneda local de Guatemala: Q 2,200.00
    """
    return f"Q {monto:,.2f}"

def obtener_nombre_mes(mes: int, en_minuscula: bool = False) -> str:
    """
    Devuelve el nombre del mes en idioma español.
    """
    meses = {
        1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
        5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
        9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
    }
    nombre = meses.get(mes, "")
    return nombre.lower() if en_minuscula else nombre

def reemplazar_texto_en_parrafo(parrafo, mapa_reemplazos: dict):
    """
    Sustituye texto dentro de un párrafo de Word. Consolida todos los runs 
    en una sola cadena para evitar fallas si Word segmentó el contenido, 
    garantizando que se mantenga el estilo, fuentes y alineación original.
    """
    texto_completo = "".join([run.text for run in parrafo.runs])
    
    modificado = False
    for buscar, reemplazar in mapa_reemplazos.items():
        if buscar in texto_completo:
            texto_completo = texto_completo.replace(buscar, reemplazar)
            modificado = True
            
    if modificado:
        if parrafo.runs:
            # Asignar todo el texto modificado al primer run para heredar estilo primario
            primer_run = parrafo.runs[0]
            primer_run.text = texto_completo
            # Vaciar los runs subsiguientes para que no repitan el texto original viejo
            for run in parrafo.runs[1:]:
                run.text = ""
        else:
            parrafo.text = texto_completo

def procesar_documento_word(ruta_plantilla: str, ruta_salida: str, mapa_reemplazos: dict):
    """
    Abre la plantilla .docx, itera sobre los párrafos normales y dentro de celdas 
    de tablas, aplica los reemplazos y guarda el nuevo documento conservando formatos.
    """
    doc = Document(ruta_plantilla)

    # 1. Procesar párrafos principales
    for parrafo in doc.paragraphs:
        reemplazar_texto_en_parrafo(parrafo, mapa_reemplazos)

    # 2. Procesar párrafos embebidos en tablas estructurales
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    reemplazar_texto_en_parrafo(parrafo, mapa_reemplazos)

    # Guardar manteniendo intactos márgenes, fuentes y tablas originales
    doc.save(ruta_salida)


# =====================================================================
# MODULE: XML READER
# =====================================================================

def leer_empleados_xml(ruta_xml: str) -> List[Dict[str, Any]]:
    """
    Lee el árbol XML de empleados y limpia el contenido de sus etiquetas.
    """
    ruta = Path(ruta_xml)
    if not ruta.exists():
        raise FileNotFoundError(f"No se encontró el archivo XML en: {ruta_xml}")

    tree = ET.parse(ruta)
    root = tree.getroot()
    
    empleados = []
    
    for emp_node in root.findall('empleado'):
        nombre = emp_node.find('nombre').text.strip() if emp_node.find('nombre') is not None else ""
        dpi = emp_node.find('dpi').text.strip() if emp_node.find('dpi') is not None else ""
        fecha_inicio_str = emp_node.find('fecha_inicio').text.strip() if emp_node.find('fecha_inicio') is not None else ""
        salario_str = emp_node.find('salario').text.strip() if emp_node.find('salario') is not None else "0.00"
        
        # Parseo de fecha al formato de recepción de la institución (DD-MES-AAAA)
        try:
            fecha_dt = datetime.strptime(fecha_inicio_str, "%Y-%m-%d")
            meses_es = ["ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO", 
                        "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE"]
            fecha_inicio = f"{fecha_dt.day:02d}-{meses_es[fecha_dt.month - 1]}-{fecha_dt.year}"
        except Exception:
            fecha_inicio = fecha_inicio_str

        empleados.append({
            "nombre": nombre,
            "dpi": dpi,
            "fecha_inicio": fecha_inicio,
            "salario": float(salario_str)
        })
        
    return empleados


# =====================================================================
# MODULE: CORE GENERATOR & CONSOLE UI
# =====================================================================

def construir_mapa_reemplazos(empleado: Dict[str, Any], anio: int, mes: int) -> dict:
    """
    Calcula todas las cadenas dependientes del mes/año para un empleado dado
    y arma el mapeo estructurado de búsqueda/reemplazo utilizado sobre la
    plantilla Word. Se aísla en su propia función para poder generar el
    recibo de cada mes del rango sin duplicar esta lógica.
    """
    ultimo_dia = obtener_ultimo_dia(anio, mes)
    nombre_mes_cap = obtener_nombre_mes(mes, en_minuscula=False)
    nombre_mes_min = obtener_nombre_mes(mes, en_minuscula=True)
    monto_formateado = formatear_moneda(empleado["salario"])

    # Formatos de cadena requeridos según reglas institucionales
    periodo_pago = f"Del 1 al {ultimo_dia} de {nombre_mes_cap} {anio}"
    percepcion = f"1 al {ultimo_dia} de {nombre_mes_min} de {anio}"
    fecha_recepcion = f"{ultimo_dia}-{mes}-{anio}"
    texto_parrafo = f"RECIBÍ DE COLEGIO MARÍA MONTESSORI LA CANTIDAD DE    {monto_formateado}  POR CONCEPTO DE SALARIO MENSUAL, SIN QUE A LA FECHA TENGA ALGÚN ADEUDO CON MI PERSONA."

    # Mapeo estructurado para evitar que colisionen fragmentos pequeños de strings
    return {
        "NOMBRE DEL TRABAJADOR:,CINTHYA JEANNETH GARCÍA MARROQUÍN": f"NOMBRE DEL TRABAJADOR:,{empleado['nombre']}",
        "CINTHYA JEANNETH GARCÍA MARROQUÍN": empleado["nombre"],
        "3737 93944 0101": empleado["dpi"],
        "3737939440101": empleado["dpi"],
        "Del 1 al 30 de Abril 2026": periodo_pago,
        "06-MAYO-2024": empleado["fecha_inicio"],
        "1 al 30 de abril de 2026": percepcion,
        "RECIBÍ DE COLEGIO MARÍA MONTESSORI LA CANTIDAD DE    Q 2,200.00  POR CONCEPTO DE SALARIO MENSUAL, SIN QUE A LA FECHA TENGA ALGÚN ADEUDO CON MI PERSONA.": texto_parrafo,
        "Q 2,200.00": monto_formateado,
        "30-4-2026": fecha_recepcion,
        "NOMBRE:,": f"NOMBRE: {empleado['nombre']}",
        "FIRMA:,": "FIRMA: "
    }


def insertar_salto_pagina_inicial(doc: Document):
    """
    Inserta, como primer elemento del cuerpo del documento, un párrafo vacío
    que contiene un salto de página duro. Se aplica sobre cada recibo
    mensual "secundario" antes de fusionarlo, para garantizar que cada mes
    comience en una hoja nueva dentro del documento consolidado del
    empleado, sin alterar encabezados, pies de página ni el contenido.
    """
    nuevo_parrafo = doc.add_paragraph()
    nuevo_parrafo_xml = nuevo_parrafo._p
    nuevo_parrafo_xml.getparent().remove(nuevo_parrafo_xml)
    doc.element.body.insert(0, nuevo_parrafo_xml)
    nuevo_parrafo.add_run().add_break(WD_BREAK.PAGE)


def combinar_recibos_mensuales(rutas_mensuales: List[Path], ruta_salida: Path):
    """
    Fusiona, en orden cronológico, los recibos mensuales individuales
    (generados previamente en una carpeta temporal a partir de la plantilla)
    en un único archivo .docx por empleado. Utiliza docxcompose.Composer,
    que concatena los documentos a nivel de las partes internas de Word
    (párrafos, tablas, imágenes, estilos), preservando exactamente el
    formato original de la plantilla en cada página. Cada recibo, a partir
    del segundo, comienza en una hoja nueva mediante un salto de página.
    """
    doc_maestro = Document(str(rutas_mensuales[0]))
    compositor = Composer(doc_maestro)

    for ruta_extra in rutas_mensuales[1:]:
        doc_secundario = Document(str(ruta_extra))
        insertar_salto_pagina_inicial(doc_secundario)
        compositor.append(doc_secundario)

    compositor.save(str(ruta_salida))


def generar_recibos(
    empleados: list, 
    ruta_plantilla: str, 
    carpeta_salida: str, 
    anio: int, 
    mes_inicial: int, 
    mes_final: int,
    callback_progreso
):
    """
    Cicla entre empleados y rango de meses, calculando los datos cambiantes
    de cada recibo de pago. Por cada empleado genera, en una carpeta
    temporal, un recibo Word individual por mes (reutilizando la plantilla
    y la lógica de reemplazo tal cual), y luego los consolida en un único
    archivo "<Nombre del empleado>.docx", en el que cada mes ocupa su
    propia página, en el orden del rango solicitado.
    """
    base_output = Path(carpeta_salida)
    base_output.mkdir(parents=True, exist_ok=True)

    for empleado in empleados:
        nombre_archivo = empleado["nombre"].replace("/", "_").replace("\\", "_").strip()
        ruta_consolidada = base_output / f"{nombre_archivo}.docx"

        # Carpeta temporal exclusiva para los recibos mensuales de este
        # empleado; se elimina automáticamente al salir del bloque "with",
        # sin dejar archivos intermedios en la carpeta de destino final.
        with tempfile.TemporaryDirectory(prefix="recibos_tmp_") as carpeta_temp:
            rutas_mensuales = []

            for mes in range(mes_inicial, mes_final + 1):
                mapa_reemplazos = construir_mapa_reemplazos(empleado, anio, mes)
                nombre_mes_cap = obtener_nombre_mes(mes, en_minuscula=False)

                archivo_temporal = Path(carpeta_temp) / f"{nombre_mes_cap}.docx"
                procesar_documento_word(ruta_plantilla, str(archivo_temporal), mapa_reemplazos)
                rutas_mensuales.append(archivo_temporal)

                # Notificar progreso a la interfaz de terminal
                callback_progreso()

            combinar_recibos_mensuales(rutas_mensuales, ruta_consolidada)

def dibujar_barra_progreso(actual: int, total: int, ancho: int = 40):
    """
    Dibuja una barra animada nativa en la consola actualizando la línea activa.
    """
    porcentaje = (actual / total) * 100
    llenado = int(ancho * actual // total)
    barra = "█" * llenado + "-" * (ancho - llenado)
    sys.stdout.write(f"\rProgreso: |{barra}| {porcentaje:.1f}% ({actual}/{total})")
    sys.stdout.flush()

def main():
    print("==================================================")
    print("  SISTEMA AUTOMATIZADO DE GENERACIÓN DE RECIBOS   ")
    print("==================================================\n")
    
    # 1. Entrada de datos por consola
    ruta_plantilla = input("1. Ruta de la plantilla Word (.docx): ").strip()
    ruta_xml = input("2. Ruta del archivo XML: ").strip()
    carpeta_salida = input("3. Carpeta donde guardar los documentos: ").strip()
    
    try:
        anio = int(input("4. Año (ej. 2026): ").strip())
        mes_inicial = int(input("5. Mes inicial (1-12): ").strip())
        mes_final = int(input("6. Mes final (1-12): ").strip())
        
        if not (1 <= mes_inicial <= 12) or not (1 <= mes_final <= 12) or mes_inicial > mes_final:
            print("[Error] El rango de meses seleccionado es inválido.")
            return
            
    except ValueError:
        print("[Error] El año y los meses ingresados deben ser números enteros numéricos.")
        return

    # 2. Validación física de la existencia de dependencias
    if not Path(ruta_plantilla).exists():
        print(f"[Error] La plantilla Word no existe en la ruta dada: {ruta_plantilla}")
        return
    if not Path(ruta_xml).exists():
        print(f"[Error] El archivo XML no existe en la ruta dada: {ruta_xml}")
        return

    print("\n[Proceso] Analizando el archivo XML de origen...")
    try:
        empleados = leer_empleados_xml(ruta_xml)
        if not empleados:
            print("[Advertencia] No se detectaron nodos de empleados en el XML.")
            return
        print(f"[Éxito] Se cargaron correctamente {len(empleados)} empleados.")
    except Exception as e:
        print(f"[Error] No se pudo leer el archivo XML: {e}")
        return

    # 3. Preparación de métricas de procesamiento
    total_meses = (mes_final - mes_inicial) + 1
    total_archivos = len(empleados) * total_meses
    print(f"[Proceso] Se generarán {total_archivos} recibos en total.\n")

    contador = 0
    dibujar_barra_progreso(contador, total_archivos)

    def actualizar_progreso():
        nonlocal contador
        contador += 1
        dibujar_barra_progreso(contador, total_archivos)

    # 4. Lanzamiento del motor automático
    try:
        generar_recibos(
            empleados=empleados,
            ruta_plantilla=ruta_plantilla,
            carpeta_salida=carpeta_salida,
            anio=anio,
            mes_inicial=mes_inicial,
            mes_final=mes_final,
            callback_progreso=actualizar_progreso
        )
        print("\n\n==================================================")
        print(" ¡Proceso completado de manera exitosa! ")
        print(f" Se generó 1 archivo .docx por empleado ({len(empleados)} en total),")
        print(f" cada uno con {total_meses} recibo(s), uno por página.")
        print(f" Destino: {Path(carpeta_salida).resolve()}")
        print("==================================================")
        
    except Exception as e:
        print(f"\n\n[Error Crítico] Falló la ejecución en lote: {e}")

if __name__ == "__main__":
    main()